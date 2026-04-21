import os
import subprocess
import sys

# Ensure python-docx is installed
subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])

# Now import 
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read the master markdown file
with open('master_document.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Split content into sections
sections = md_content.split('## ')
for idx, section in enumerate(sections):
    if idx == 0:
        # Title page
        title_lines = section.strip().split('\n')
        title = doc.add_heading(title_lines[0].replace('# ', ''), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in title_lines[1:]:
            if line.strip():
                p = doc.add_paragraph(line.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Process sections
        lines = section.split('\n')
        section_title = lines[0].strip()
        
        # Add section heading
        heading = doc.add_heading(section_title, level=1)
        
        # Process content
        content = '\n'.join(lines[1:])
        paragraphs = content.split('\n\n')
        
        for para in paragraphs:
            para = para.strip()
            if para.startswith('###'):
                # Subheading
                doc.add_heading(para.replace('###', '').strip(), level=2)
            elif para.startswith('- '):
                # Bullet points
                p = doc.add_paragraph(para[2:], style='List Bullet')
            elif para.startswith('**') and para.endswith('**'):
                # Bold text
                p = doc.add_paragraph()
                p.add_run(para[2:-2]).bold = True
            elif para.startswith('*') and para.endswith('*'):
                # Italic text
                p = doc.add_paragraph()
                p.add_run(para[1:-1]).italic = True
            elif 'References' in para:
                # References section
                doc.add_heading('References', level=1)
                ref_lines = para.split('\n')[1:]
                for ref in ref_lines:
                    if ref.strip() and not ref.startswith('---'):
                        p = doc.add_paragraph(ref.strip(), style='List Number')
            else:
                # Regular paragraph
                if para.strip():
                    doc.add_paragraph(para.strip())

# Save as DOCX
doc.save('Zimbabwe_Parliament_Intervention_Notes_2025.docx')
print("DOCX file created successfully: Zimbabwe_Parliament_Intervention_Notes_2025.docx")