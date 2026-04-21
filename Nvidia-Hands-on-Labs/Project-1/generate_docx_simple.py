
import os
import sys
import subprocess

# Install docx using ensurepip
subprocess.check_call([sys.executable, '-m', 'ensurepip', '--upgrade'])
subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])

# Now import
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read the master markdown file
with open('master_document.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Split content into lines
lines = md_content.split('\n')
current_paragraph = None

for line in lines:
    line = line.strip()
    if not line:
        continue
    
    if line.startswith('# '):
        # Main title
        p = doc.add_heading(line[2:], level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith('## '):
        # Section heading
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        # Subsection heading
        doc.add_heading(line[4:], level=2)
    elif line.startswith('- '):
        # Bullet point
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('**') and line.endswith('**'):
        # Bold text
        p = doc.add_paragraph()
        p.add_run(line[2:-2]).bold = True
    elif line.startswith('*') and line.endswith('*'):
        # Italic text
        p = doc.add_paragraph()
        p.add_run(line[1:-1]).italic = True
    elif 'References' in line and line.startswith('##'):
        # References section
        doc.add_heading('References', level=1)
    else:
        # Regular paragraph
        if line and not line.startswith('---'):
            doc.add_paragraph(line)

# Save as DOCX
doc.save('Zimbabwe_Parliament_Intervention_Notes_2025.docx')
print("DOCX file created successfully: Zimbabwe_Parliament_Intervention_Notes_2025.docx")